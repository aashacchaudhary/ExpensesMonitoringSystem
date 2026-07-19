from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT = BASE_DIR / "documentation" / "Expense_Management_Code_Documentation.docx"

ACCENT = "16A085"
NAVY = "1E2A38"
MUTED = "6B7280"
BORDER = "D9E2EC"
LIGHT = "F6F8FB"

CODE_FILES = [
    (
        "requirements.txt",
        "Lists the Python dependency needed to run the website. Installing this file downloads Django.",
    ),
    (
        "manage.py",
        "Main Django command runner. It connects terminal commands to the project settings.",
    ),
    (
        "expense_manager/__init__.py",
        "Marks the expense_manager folder as a Python package.",
    ),
    (
        "expense_manager/settings.py",
        "Stores global Django configuration such as installed apps, middleware, database, templates, static files, and authentication redirects.",
    ),
    (
        "expense_manager/urls.py",
        "Root URL router. It connects Django admin and includes all expenses app routes.",
    ),
    (
        "expense_manager/asgi.py",
        "ASGI deployment entry point for async-capable Python servers.",
    ),
    (
        "expense_manager/wsgi.py",
        "WSGI deployment entry point for standard Django hosting servers.",
    ),
    (
        "expenses/__init__.py",
        "Marks the expenses folder as a Python package.",
    ),
    (
        "expenses/apps.py",
        "Defines the Django app configuration for the expenses application.",
    ),
    (
        "expenses/models.py",
        "Defines the database tables and relationships for expenses, budgets, family groups, and family memberships.",
    ),
    (
        "expenses/forms.py",
        "Defines Django forms used for registration, login, adding expenses, filtering statements, reports, budgets, and family group workflows.",
    ),
    (
        "expenses/views.py",
        "Contains the main website logic. Views read requests, query models, validate forms, generate reports, and render templates or downloads.",
    ),
    (
        "expenses/urls.py",
        "Maps app URL paths to view functions and stable URL names used in templates.",
    ),
    (
        "expenses/admin.py",
        "Registers app models in Django admin and configures list display, filters, and search fields.",
    ),
    (
        "expenses/tests.py",
        "Automated tests for registration, privacy, budget calculations, date filtering, report downloads, and authenticated page rendering.",
    ),
    (
        "expenses/migrations/0001_initial.py",
        "Initial migration generated from models.py. It creates the database tables for the expenses app.",
    ),
    (
        "expenses/migrations/__init__.py",
        "Marks the migrations folder as a Python package so Django can discover migration files.",
    ),
    (
        "expenses/templates/expenses/base.html",
        "Shared base template. It defines the Soft Ledger Glass layout, navigation, messages, CSS, scripts, and the content block used by all pages.",
    ),
    (
        "expenses/templates/expenses/home.html",
        "Public landing page shown before login.",
    ),
    (
        "expenses/templates/expenses/login.html",
        "Login page template that renders Django's authentication form with CSRF protection and error messages.",
    ),
    (
        "expenses/templates/expenses/register.html",
        "Registration page template that renders the account creation form and validation errors.",
    ),
    (
        "expenses/templates/expenses/dashboard.html",
        "Dashboard page showing totals, budget mood ring, charts, budget alerts, and recent transactions.",
    ),
    (
        "expenses/templates/expenses/expense_form.html",
        "Add expense page template that renders ExpenseForm and posts the new expense to the backend.",
    ),
    (
        "expenses/templates/expenses/expense_list.html",
        "Expense statement page with search, category/payment/date filters, custom date range, desktop table, and mobile cards.",
    ),
    (
        "expenses/templates/expenses/monthly_report.html",
        "My Reports page with month/year filter, charts, summaries, statement rows, and PDF/CSV download buttons.",
    ),
    (
        "expenses/templates/expenses/budget_form.html",
        "Budget setup page for saving a monthly budget amount.",
    ),
    (
        "expenses/templates/expenses/family_group.html",
        "Family Reports management page for creating and joining family groups.",
    ),
    (
        "expenses/templates/expenses/family_report.html",
        "Family Reports page with combined household totals, charts, statement rows, and PDF/CSV downloads.",
    ),
    (
        "expenses/templates/expenses/profile.html",
        "Profile page showing account information and user activity metrics.",
    ),
    (
        "expenses/static/expenses/css/styles.css",
        "Custom Soft Ledger Glass stylesheet for colors, layout, cards, sidebar, forms, tables, badges, charts, and mobile responsiveness.",
    ),
    (
        "expenses/static/expenses/js/charts.js",
        "JavaScript helper that reads Django JSON chart data and renders Chart.js doughnut, bar, and line charts.",
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    dxa_widths = [int(width * 1440) for width in widths]
    table_width = sum(dxa_widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in dxa_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(0, tbl_grid)

    for row in table.rows:
        for idx, width in enumerate(dxa_widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_header_footer(document):
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Expense Management System - Code Documentation"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(107, 114, 128)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = "Generated documentation"
    run = paragraph.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(107, 114, 128)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, NAVY),
        ("Subtitle", 11, MUTED),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, NAVY),
        ("Heading 3", 11, NAVY),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if "Heading" in name or name == "Title":
            style.font.bold = True


def add_title(document):
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Expense Management System\nCode Documentation")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "A file-by-file and workflow-level explanation of the Django expense monitoring website."
    )
    meta = document.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_table_width(meta, [2.0, 4.5])
    data = [
        ("Project", "Django Expense Management System"),
        ("Theme", "Soft Ledger Glass"),
        ("Generated On", date.today().strftime("%B %d, %Y")),
        ("Database", "SQLite through Django ORM"),
    ]
    for row, (label, value) in zip(meta.rows, data):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
    document.add_paragraph()


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "EAF4F1")
        paragraph = cell.paragraphs[0]
        paragraph.runs[0].bold = True
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = str(value)
    set_table_width(table, widths)
    document.add_paragraph()
    return table


def add_code_block(document, code):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_code_listing(document, path, description):
    source_path = BASE_DIR / path
    document.add_heading(path, level=2)
    document.add_paragraph(description)
    if not source_path.exists():
        document.add_paragraph("File was not found in the project workspace.")
        return

    code = source_path.read_text(encoding="utf-8")
    if code == "":
        code = "# Empty file"

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string("202124")


def add_code_appendix(document):
    document.add_heading("15. Complete Code Listings and Explanations", level=1)
    document.add_paragraph(
        "This section contains the source code used by the website. Each listing starts with a short explanation "
        "of what the file does, followed by the actual code from the project."
    )
    for path, description in CODE_FILES:
        add_code_listing(document, path, description)


def add_architecture(document):
    document.add_heading("1. Project Overview", level=1)
    document.add_paragraph(
        "This website is a Django application for personal and family expense monitoring. "
        "It uses Django authentication for users, SQLite for persistent storage, Django templates "
        "for server-rendered pages, and Chart.js for visual reports."
    )
    add_bullets(
        document,
        [
            "Users register, log in, and manage their own expenses.",
            "Each expense is saved with title, amount, category, payment method, date, and description.",
            "Budgets are stored per user, month, and year, then compared against monthly spending.",
            "Family groups allow multiple users to combine monthly expense reports safely.",
            "CSV and PDF report downloads are generated by the report views from live database data.",
        ],
    )

    document.add_heading("2. Technology Stack", level=1)
    add_table(
        document,
        ["Layer", "Technology", "Purpose"],
        [
            ("Backend", "Python + Django", "Handles routing, authentication, views, ORM queries, and report generation."),
            ("Database", "SQLite", "Stores users, expenses, budgets, family groups, memberships, sessions, and admin data."),
            ("Frontend", "Django Templates, HTML, CSS, JavaScript", "Renders dynamic pages using backend context variables."),
            ("UI Framework", "Bootstrap 5 + Bootstrap Icons", "Provides responsive grid, forms, icons, buttons, and layout helpers."),
            ("Charts", "Chart.js", "Draws category, monthly, daily, and family report graphs from JSON data."),
            ("Theme", "Soft Ledger Glass CSS", "Custom glass-style sidebar, cards, tables, badges, forms, and mobile navigation."),
        ],
        [1.25, 1.7, 3.55],
    )

    document.add_heading("3. How a Request Is Processed", level=1)
    add_numbered(
        document,
        [
            "The browser requests a URL such as /expenses/ or /reports/monthly/.",
            "expense_manager/urls.py sends all normal site routes into expenses/urls.py.",
            "expenses/urls.py maps the URL name to the correct view function or class-based view.",
            "The view checks authentication where needed, reads form/query data, and queries the database through models.",
            "The view builds a context dictionary and renders a Django template.",
            "The template extends base.html, displays values, and loads static CSS and JavaScript.",
            "For charts, views send JSON-safe labels and values; charts.js turns that data into Chart.js graphs.",
            "For CSV/PDF downloads, the report view returns an HttpResponse attachment instead of rendering HTML.",
        ],
    )


def add_project_files(document):
    document.add_heading("4. Project Structure and File Roles", level=1)
    add_table(
        document,
        ["Path", "Role"],
        [
            ("manage.py", "Command-line launcher for Django tasks such as runserver, migrate, makemigrations, createsuperuser, and test."),
            ("requirements.txt", "Lists Python dependencies. This project requires Django 5.x."),
            ("expense_manager/settings.py", "Global settings for installed apps, middleware, templates, database, static files, auth redirects, and timezone."),
            ("expense_manager/urls.py", "Root URL file. It connects Django admin and includes the expenses app routes."),
            ("expense_manager/asgi.py and wsgi.py", "Deployment entry points for ASGI and WSGI servers."),
            ("expenses/models.py", "Database schema for Expense, Budget, FamilyGroup, and FamilyMembership."),
            ("expenses/forms.py", "Form classes for registration, login, expense entry, searching, reports, budgets, and family joining."),
            ("expenses/views.py", "Main business logic: dashboard, expenses, reports, downloads, budget alerts, family flows, and profile."),
            ("expenses/urls.py", "Named app routes used by templates and redirects."),
            ("expenses/admin.py", "Registers models in Django admin with filtering and search support."),
            ("expenses/templates/expenses/*.html", "Django templates for all pages and the shared base layout."),
            ("expenses/static/expenses/css/styles.css", "Soft Ledger Glass visual theme."),
            ("expenses/static/expenses/js/charts.js", "Chart.js helper functions and chart styling."),
            ("expenses/tests.py", "Automated tests for auth, privacy, reports, downloads, filters, and page rendering."),
        ],
        [2.5, 4.0],
    )


def add_settings_urls(document):
    document.add_heading("5. Core Django Configuration", level=1)
    document.add_heading("5.1 manage.py", level=2)
    document.add_paragraph(
        "manage.py sets DJANGO_SETTINGS_MODULE to expense_manager.settings and calls Django's "
        "execute_from_command_line. This is why commands like python manage.py runserver and "
        "python manage.py migrate know which project settings to use."
    )
    add_code_block(document, 'os.environ.setdefault("DJANGO_SETTINGS_MODULE", "expense_manager.settings")')

    document.add_heading("5.2 settings.py", level=2)
    add_bullets(
        document,
        [
            "INSTALLED_APPS enables Django admin, authentication, sessions, messages, static files, and the expenses app.",
            "MIDDLEWARE enables security, sessions, CSRF protection, authentication, messages, and clickjacking protection.",
            "DATABASES configures SQLite at BASE_DIR / db.sqlite3.",
            "TEMPLATES enables APP_DIRS so Django can find templates inside expenses/templates/.",
            "LOGIN_URL, LOGIN_REDIRECT_URL, and LOGOUT_REDIRECT_URL control auth redirects.",
            "TIME_ZONE is set to Asia/Kathmandu and USE_TZ is enabled.",
        ],
    )

    document.add_heading("5.3 URL Configuration", level=2)
    document.add_paragraph(
        "expense_manager/urls.py includes the app routes at the site root. expenses/urls.py "
        "then maps each page to a view and gives it a stable URL name used throughout templates."
    )
    add_table(
        document,
        ["URL", "View", "URL Name"],
        [
            ("/", "home", "home"),
            ("/register/", "register", "register"),
            ("/login/", "ExpenseLoginView", "login"),
            ("/dashboard/", "dashboard", "dashboard"),
            ("/expenses/add/", "expense_create", "expense_create"),
            ("/expenses/", "expense_list", "expense_list"),
            ("/reports/monthly/", "monthly_report", "monthly_report"),
            ("/family/report/", "family_report", "family_report"),
        ],
        [1.7, 2.2, 2.1],
    )


def add_models(document):
    document.add_heading("6. Database Models", level=1)
    document.add_paragraph(
        "The app uses Django's built-in User model for authentication and links every user-owned record "
        "back to that user with ForeignKey relationships."
    )
    add_table(
        document,
        ["Model", "Important Fields", "What It Does"],
        [
            ("Expense", "user, title, amount, category, date, payment_method, description, created_at, updated_at", "Stores one spending record. Queries filter by logged-in user so users only see their own data."),
            ("Budget", "user, month, year, amount", "Stores one monthly budget per user/month/year. A unique constraint prevents duplicate budgets for the same period."),
            ("FamilyGroup", "name, code, created_by, created_at", "Stores a family group. save() generates a unique join code when the group is first created."),
            ("FamilyMembership", "user, family_group, joined_at", "Connects users to family groups. A unique constraint prevents duplicate membership rows."),
        ],
        [1.4, 2.6, 2.5],
    )

    document.add_heading("6.1 Expense Model Behavior", level=2)
    add_bullets(
        document,
        [
            "category and payment_method use fixed choices so forms and filters stay consistent.",
            "amount uses DecimalField with MinValueValidator(0.01) to prevent zero or negative expenses.",
            "Meta ordering sorts newest expenses first by date and creation time.",
            "__str__ returns a readable title and amount for admin and debugging.",
        ],
    )

    document.add_heading("6.2 Family Code Generation", level=2)
    document.add_paragraph(
        "FamilyGroup.generate_unique_code builds an 8-character uppercase/digit code using Python's "
        "secrets module. It checks the database before accepting a code, preventing collisions."
    )


def add_forms(document):
    document.add_heading("7. Forms", level=1)
    add_table(
        document,
        ["Form", "Purpose", "Important Behavior"],
        [
            ("RegisterForm", "Creates a user account.", "Validates duplicate username and duplicate email, then uses Django's UserCreationForm password validation."),
            ("StyledAuthenticationForm", "Login form.", "Extends Django AuthenticationForm and applies Bootstrap form-control classes."),
            ("ExpenseForm", "Add expense form.", "Uses ModelForm fields and HTML date picker for the date field."),
            ("ExpenseSearchForm", "Expense statement search/filter form.", "Filters title, category, payment method, exact date, custom date from/to, month, and year."),
            ("ReportFilterForm", "Monthly and family report period selector.", "Defaults month and year to the current date."),
            ("BudgetForm", "Monthly budget form.", "Updates or creates one budget amount for a selected month/year."),
            ("FamilyGroupForm", "Create family group.", "Accepts the family group name."),
            ("JoinFamilyForm", "Join family group.", "Cleans family code by trimming and uppercasing it."),
        ],
        [1.55, 2.0, 2.95],
    )
    document.add_paragraph(
        "apply_bootstrap_widgets is a helper that adds form-control or form-select CSS classes to every field, "
        "keeping form styling consistent across pages."
    )


def add_views(document):
    document.add_heading("8. Views and Business Logic", level=1)
    add_table(
        document,
        ["View / Helper", "What It Does", "How It Performs the Task"],
        [
            ("home", "Shows public landing page.", "Redirects authenticated users to dashboard; otherwise renders home.html."),
            ("register", "Creates new users.", "Processes RegisterForm, saves a valid user, logs the user in, and redirects to dashboard."),
            ("ExpenseLoginView", "Logs users in.", "Uses Django LoginView with a styled auth form and invalid-login message."),
            ("dashboard", "Shows monthly summary.", "Filters current user's current-month expenses, sums totals, calculates budget balance, builds chart payloads, and sends mood ring data."),
            ("expense_create", "Adds an expense.", "Saves ExpenseForm with request.user attached, then redirects to expense_list."),
            ("expense_list", "Shows searchable statement.", "Filters the logged-in user's expenses by title, category, payment, date, date range, month, and year."),
            ("expense_delete", "Deletes an expense.", "Uses get_object_or_404 with pk and user, so users cannot delete another user's expense."),
            ("monthly_report", "Shows and downloads My Reports.", "Filters by month/year, calculates totals and summaries, renders HTML or returns CSV/PDF attachments."),
            ("budget_setup", "Saves monthly budget.", "Uses update_or_create to avoid duplicate monthly budgets."),
            ("family_group", "Creates or joins family groups.", "Uses hidden action values to distinguish create and join forms in one page."),
            ("family_report", "Shows and downloads Family Reports.", "Verifies membership, combines expenses for group members only, then renders or downloads statements."),
            ("profile", "Shows account metrics.", "Sums all-time expense, current month total, budget count, and family group count."),
        ],
        [1.45, 1.85, 3.2],
    )

    document.add_heading("8.1 Report Download Helpers", level=2)
    add_bullets(
        document,
        [
            "csv_response creates a text/csv HttpResponse and writes rows with Python's csv module.",
            "build_simple_pdf manually creates a lightweight PDF with plain text report lines.",
            "pdf_response wraps generated PDF bytes in an application/pdf attachment response.",
            "money formats Decimal amounts with two decimal places for consistent statements.",
        ],
    )

    document.add_heading("8.2 Budget Alerts and Mood Ring", level=2)
    add_bullets(
        document,
        [
            "budget_alerts compares monthly expense total against the saved budget.",
            "It returns calm alert messages for 75 percent, 90 percent, and exceeded budget states.",
            "budget_mood returns used_percent, ring_percent, label, and state for the dashboard's Spending Mood Ring.",
            "The template uses these values to color the circular progress indicator.",
        ],
    )


def add_templates_static(document):
    document.add_heading("9. Templates and Frontend", level=1)
    add_table(
        document,
        ["Template", "Purpose"],
        [
            ("base.html", "Shared layout. Contains Soft Ledger sidebar, mobile bottom nav, message alerts, CDN links, static CSS/JS, and the content block."),
            ("home.html", "Public landing page with hero section and feature cards."),
            ("login.html / register.html", "Authentication cards that keep CSRF tokens and Django form errors intact."),
            ("dashboard.html", "Summary cards, budget mood ring, category chart, monthly chart, and recent transaction list."),
            ("expense_form.html", "Add expense form with CSRF token and all model fields."),
            ("expense_list.html", "Search/filter statement page with desktop table and mobile cards."),
            ("monthly_report.html", "My Reports page with filters, charts, category summary, statement list, and PDF/CSV buttons."),
            ("family_group.html", "Family Reports management page for creating/joining groups."),
            ("family_report.html", "Family Reports dashboard with combined member/category/daily summaries and downloads."),
            ("profile.html", "Account and activity overview."),
        ],
        [2.25, 4.25],
    )

    document.add_heading("9.1 Soft Ledger Glass CSS", level=2)
    add_bullets(
        document,
        [
            "The :root variables define the main colors, radius scale, and shadows.",
            "The desktop layout uses a fixed navy sidebar; mobile uses a bottom navigation bar.",
            "Glass cards use translucent white backgrounds, borders, blur, and soft shadows.",
            "Category chips use pastel backgrounds and category-specific icons.",
            "Forms use rounded fields, teal focus states, comfortable padding, and responsive grid behavior.",
            "Tables are hidden on mobile where stacked transaction cards are shown instead.",
        ],
    )

    document.add_heading("9.2 Chart JavaScript", level=2)
    add_bullets(
        document,
        [
            "charts.js reads JSON data from Django's json_script output.",
            "renderPie creates doughnut charts for category summaries.",
            "renderBar creates rounded bar charts for monthly and member-wise totals.",
            "renderLine creates filled line charts for daily expense movement.",
            "If all values are zero or missing, the script replaces the canvas with a friendly empty state.",
        ],
    )


def add_security_and_tests(document):
    document.add_heading("10. Security and Data Privacy", level=1)
    add_bullets(
        document,
        [
            "login_required protects dashboard, expenses, budgets, reports, family, and profile pages.",
            "CSRF tokens are included in all POST forms.",
            "Expense queries always filter by request.user for personal pages.",
            "Expense deletion uses both pk and user lookup to prevent deleting another user's data.",
            "Family reports verify the user is a member of the selected family group before showing combined data.",
            "Django authentication handles password hashing, sessions, login, and logout.",
        ],
    )

    document.add_heading("11. Tests", level=1)
    add_table(
        document,
        ["Test Area", "What It Confirms"],
        [
            ("Registration", "Duplicate email validation is shown to users."),
            ("Authentication", "Dashboard redirects anonymous visitors to login."),
            ("Expense Privacy", "Users only see their own expense records."),
            ("Budget", "Remaining budget is calculated and displayed correctly."),
            ("Family Privacy", "Family report includes group members but excludes outsiders."),
            ("Page Rendering", "Authenticated pages render with status 200."),
            ("Custom Date Range", "Expense statement filters from and to calendar dates."),
            ("My Reports Downloads", "Monthly report CSV and PDF attachments are generated."),
            ("Family Reports Downloads", "Family report CSV and PDF attachments are generated."),
        ],
        [2.0, 4.5],
    )


def add_setup_and_operations(document):
    document.add_heading("12. Installation and Local Hosting", level=1)
    document.add_paragraph("On a new device, the normal setup sequence is:")
    add_code_block(
        document,
        "python -m venv venv\n"
        "venv\\Scripts\\activate\n"
        "pip install -r requirements.txt\n"
        "python manage.py migrate\n"
        "python manage.py createsuperuser\n"
        "python manage.py runserver"
    )
    add_bullets(
        document,
        [
            "pip install downloads Django from requirements.txt.",
            "migrate creates the database tables from migration files.",
            "createsuperuser is only required for admin panel access.",
            "runserver starts the local development server, usually at http://127.0.0.1:8000/.",
        ],
    )

    document.add_heading("13. End-to-End Feature Flow Examples", level=1)
    add_table(
        document,
        ["Feature", "Flow"],
        [
            ("Add Expense", "User opens /expenses/add/, submits ExpenseForm, view attaches request.user, saves Expense, then redirects to /expenses/."),
            ("Search Statement", "User opens /expenses/ with query parameters; ExpenseSearchForm validates filters; the view filters the QuerySet before rendering."),
            ("Set Budget", "User opens /budget/, submits BudgetForm; update_or_create saves one budget for that user/month/year."),
            ("My Reports Download", "User clicks PDF or CSV; monthly_report sees download parameter and returns a file attachment instead of HTML."),
            ("Family Reports", "User joins a group by code; report view verifies membership and combines expenses from all group members."),
        ],
        [1.8, 4.7],
    )


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    configure_styles(document)
    add_header_footer(document)
    add_title(document)
    add_architecture(document)
    add_project_files(document)
    add_settings_urls(document)
    add_models(document)
    add_forms(document)
    add_views(document)
    add_templates_static(document)
    add_security_and_tests(document)
    add_setup_and_operations(document)

    document.add_heading("14. Maintenance Notes", level=1)
    add_bullets(
        document,
        [
            "When adding model fields, create and apply migrations.",
            "When adding pages, define a view, URL name, template, and tests.",
            "Keep URL names stable because templates and redirects depend on them.",
            "Keep user filters in every personal-data query to protect privacy.",
            "Run python manage.py test after backend or template changes.",
        ],
    )
    add_code_appendix(document)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
